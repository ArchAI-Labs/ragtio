"""Modulo di ingestione documenti."""

import logging
from pathlib import Path
from typing import List, Optional

import pandas as pd
from haystack import Document

logger = logging.getLogger(__name__)


class UnsupportedFormatError(Exception):
    """Sollevata quando il formato del file non è supportato dal modulo di ingestione."""

    SUPPORTED = [".csv", ".json", ".jsonl", ".md", ".pdf", ".txt", ".docx"]

    def __init__(self, extension: str):
        super().__init__(
            f"Formato '{extension}' non supportato. "
            f"Formati supportati: {self.SUPPORTED}"
        )


def ingest(
    file_path: str | Path,
    content_column: Optional[str] = None,
    metadata_columns: Optional[List[str]] = None,
    extra_metadata: Optional[dict] = None,
) -> List[Document]:
    """
    Legge un file e produce una lista di Document Haystack.

    Args:
        file_path: Percorso assoluto o relativo al file da ingerire.
        content_column: Nome della colonna contenente il testo principale.
                        Obbligatorio per CSV e JSON tabulare.
                        Ignorato per PDF, TXT, DOCX.
        metadata_columns: Lista di nomi di colonne da includere come metadati.
                          Se None, vengono incluse tutte le colonne tranne content_column.
                          Ignorato per PDF, TXT, DOCX.
        extra_metadata: Dizionario di metadati aggiuntivi da aggiungere a tutti
                        i documenti prodotti (es. {"category": "medico"}).

    Returns:
        Lista di Document Haystack con content e meta popolati.

    Raises:
        FileNotFoundError: il file non esiste al percorso indicato.
        ValueError: content_column non trovata nel file CSV/JSON.
        UnicodeDecodeError: il file TXT non può essere letto (UTF-8 o latin-1).
        UnsupportedFormatError: estensione non riconosciuta.
    """
    import time

    file_path = Path(file_path)
    start = time.monotonic()
    logger.info("Inizio ingestione: %s", file_path)

    if not file_path.exists():
        raise FileNotFoundError(f"File non trovato: {file_path}")

    extra_metadata = extra_metadata or {}
    ext = file_path.suffix.lower()

    if ext == ".csv":
        documents = _ingest_csv(file_path, content_column, metadata_columns, extra_metadata)
    elif ext in (".json", ".jsonl"):
        documents = _ingest_json(file_path, content_column, metadata_columns, extra_metadata)
    elif ext == ".pdf":
        documents = _ingest_pdf(file_path, extra_metadata)
    elif ext in (".txt", ".md"):
        documents = _ingest_txt(file_path, extra_metadata)
    elif ext == ".docx":
        documents = _ingest_docx(file_path, extra_metadata)
    else:
        raise UnsupportedFormatError(ext)

    elapsed = time.monotonic() - start
    logger.info(
        "Ingestione completata: %s — %d documenti prodotti in %.2f s",
        file_path,
        len(documents),
        elapsed,
    )
    return documents


def _ingest_csv_like(
    df: pd.DataFrame,
    file_path: Path,
    content_column: str,
    metadata_columns: Optional[List[str]],
    extra_metadata: dict,
) -> List[Document]:
    """Logica comune per CSV e JSON tabulare."""
    if content_column not in df.columns:
        raise ValueError(
            f"Colonna '{content_column}' non trovata nel file. "
            f"Colonne disponibili: {list(df.columns)}"
        )
    meta_cols = metadata_columns or [c for c in df.columns if c != content_column]
    documents = []
    for idx, row in df.iterrows():
        content = str(row[content_column])
        if not content.strip():
            logger.warning("Riga %d vuota in %s, saltata.", idx, file_path)
            continue
        meta = {col: row[col] for col in meta_cols if col in row and pd.notna(row[col])}
        meta["source"] = str(file_path)
        meta["row_index"] = int(idx)
        meta["id"] = f"{file_path.stem}_row_{idx}"
        meta.update(extra_metadata)
        documents.append(Document(content=content, meta=meta))
    return documents


def _ingest_csv(
    file_path: Path,
    content_column: str,
    metadata_columns: Optional[List[str]],
    extra_metadata: dict,
) -> List[Document]:
    df = pd.read_csv(file_path, encoding="utf-8")
    return _ingest_csv_like(df, file_path, content_column, metadata_columns, extra_metadata)


def _ingest_json(
    file_path: Path,
    content_column: str,
    metadata_columns: Optional[List[str]],
    extra_metadata: dict,
) -> List[Document]:
    if file_path.suffix == ".jsonl":
        df = pd.read_json(file_path, lines=True)
    else:
        df = pd.read_json(file_path)
    return _ingest_csv_like(df, file_path, content_column, metadata_columns, extra_metadata)


def _ingest_pdf(file_path: Path, extra_metadata: dict) -> List[Document]:
    import fitz  # PyMuPDF

    doc = fitz.open(str(file_path))
    documents = []
    for page_num, page in enumerate(doc):
        text = page.get_text("text").strip()
        if not text:
            logger.warning("Pagina %d vuota in %s, saltata.", page_num, file_path)
            continue
        meta = {
            "source": str(file_path),
            "id": f"{file_path.stem}_page_{page_num}",
            "page": page_num + 1,
        }
        meta.update(extra_metadata)
        documents.append(Document(content=text, meta=meta))
    return documents


def _ingest_txt(file_path: Path, extra_metadata: dict) -> List[Document]:
    try:
        text = file_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        text = file_path.read_text(encoding="latin-1")
    meta = {
        "source": str(file_path),
        "id": file_path.stem,
    }
    meta.update(extra_metadata)
    return [Document(content=text.strip(), meta=meta)]


def _ingest_docx(file_path: Path, extra_metadata: dict) -> List[Document]:
    from docx import Document as DocxDocument

    doc = DocxDocument(str(file_path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    full_text = "\n\n".join(paragraphs)
    meta = {
        "source": str(file_path),
        "id": file_path.stem,
    }
    meta.update(extra_metadata)
    return [Document(content=full_text, meta=meta)]
